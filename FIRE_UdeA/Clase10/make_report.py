import os
import glob
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches, Pt

try:
    from FIRE_UdeA_model_v2 import engineer_features, get_feature_columns
except ImportError:
    engineer_features = None
    get_feature_columns = None

sns.set(style="whitegrid", context="talk")

def error_analysis_plot():
    # Load and process data to find clustering errors
    script_dir = os.path.dirname(os.path.abspath(__file__))
    csv_path = os.path.join(script_dir, "dataset_sintetico_FIRE_UdeA_realista.csv")
    
    if not os.path.exists(csv_path):
        return None, "Dataset no encontrado."
        
    origin_df = pd.read_csv(csv_path)
    df_processed = origin_df.copy()
    
    if engineer_features is not None:
        try:
            df_processed = engineer_features(df_processed)
            feature_cols = get_feature_columns(df_processed)
        except Exception:
            feature_cols = [c for c in df_processed.columns if c not in ['anio', 'unidad', 'label']]
    else:
        feature_cols = [c for c in df_processed.columns if c not in ['anio', 'unidad', 'label']]
        
    for col in feature_cols:
        if df_processed[col].isnull().any():
            df_processed[col] = df_processed[col].fillna(df_processed[col].median())
            
    X = StandardScaler().fit_transform(df_processed[feature_cols].values)
    
    # K-Means k=2
    kmeans = KMeans(n_clusters=2, random_state=42, n_init='auto')
    preds = kmeans.fit_predict(X)
    
    # True labels
    if 'label' not in df_processed.columns:
        return None, "Columna 'label' no existe."
    real = df_processed['label'].values
    
    # Align labels
    acc1 = np.mean(preds == real)
    if acc1 < 0.5:
        preds = 1 - preds
        
    df_processed['pred'] = preds
    # An error is when the boolean does not match
    df_processed['error'] = df_processed['label'] != df_processed['pred']
    
    errores_f = df_processed.groupby('unidad')['error'].sum().reset_index()
    errores_f = errores_f.sort_values('error', ascending=False)
    
    peor_facultad = errores_f.iloc[0]['unidad']
    max_errores = errores_f.iloc[0]['error']
    total_errores = errores_f['error'].sum()
    print(f"\n[ANALYSIS RESULT] La facultad que más se equivocó: '{peor_facultad}' con {max_errores} errores.\n")
    
    # Plot
    plt.figure(figsize=(10, 6))
    sns.barplot(data=errores_f, x='error', y='unidad', palette='Reds_r')
    plt.title('Errores de Clasificación (K-Means) por Facultad')
    plt.xlabel('Cantidad de Errores')
    plt.ylabel('Facultad')
    plt.tight_layout()
    plot_path = os.path.join(script_dir, "errores_por_facultad.png")
    plt.savefig(plot_path)
    plt.close()
    
    summary_text = (f"El algoritmo K-Means cometió {total_errores} errores en total al compararlo " 
                    f"con la variable real de riesgo. La unidad/facultad que más se equivocó fue "
                    f"'{peor_facultad}' con un total de {max_errores} clasificaciones incorrectas.")
    return plot_path, summary_text


def main():
    doc = Document()
    
    # Title
    doc.add_heading('Reporte de Clustering - Clase 10', 0)
    
    p = doc.add_paragraph()
    p.add_run('Autor: ').bold = True
    p.add_run('Agustín Figueroa Sierra\n')
    
    # Error Analysis Section
    doc.add_heading('1. Análisis de Errores por Facultad', level=1)
    
    doc.add_paragraph('Para identificar qué facultad se agrupó de forma más errónea, utilizamos los resultados del clustering de K-Means (con 2 grupos) mapeándolos y validándolos contra la variable dependiente de riesgo ("label"). Para esto calculamos la diferencia entre la etiqueta real y el clúster asignado a cada objeto evaluado.')
    
    plot_path, summary_text = error_analysis_plot()
    if plot_path and os.path.exists(plot_path):
        doc.add_paragraph(summary_text)
        try:
            doc.add_picture(plot_path, width=Inches(6.0))
        except Exception as e:
            doc.add_paragraph(f"[No se pudo insertar la imagen: {e}]")
    else:
        doc.add_paragraph(f"Error en el análisis: {summary_text}")
        
    doc.add_paragraph()
    
    # General Results Section
    doc.add_heading('2. Resultados de Algoritmos (DBSCAN, HDBSCAN, Subtractive, FCM, KMeans)', level=1)
    
    results_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'results', 'clustering')
    png_files = glob.glob(os.path.join(results_dir, '**', '*.png'), recursive=True)
    
    datasets = {}
    for path in png_files:
        dataset_name = os.path.basename(os.path.dirname(path))
        if dataset_name == 'clustering': dataset_name = 'General'
        if dataset_name not in datasets: datasets[dataset_name] = []
        datasets[dataset_name].append(path)
        
    # Ordenar datasets
    def get_sort_key(name):
        return name
        
    for ds_name in sorted(datasets.keys(), key=get_sort_key):
        doc.add_heading(f'Dataset: {ds_name}', level=2)
        
        # Sort files so related stay together
        def sort_files(f):
            return os.path.basename(f)
            
        for img_path in sorted(datasets[ds_name], key=sort_files):
            img_name = os.path.basename(img_path).replace('.png', '').replace('_', ' ').title()
            doc.add_heading(img_name, level=3)
            try:
                doc.add_picture(img_path, width=Inches(6.0))
            except Exception as e:
                doc.add_paragraph(f"[No se pudo insertar la imagen {img_name}: {e}]")
            doc.add_paragraph()
            
    out_name = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Figueroa_Sierra_Agustin-Clase10.docx')
    doc.save(out_name)
    print(f"Reporte generado exitosamente: {out_name}")

if __name__ == '__main__':
    main()
